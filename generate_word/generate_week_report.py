from docx import Document
from docx.oxml.ns import qn
from datetime import datetime, timedelta
from fire import Fire
import sys


def get_monday_sunday(delta=0):
    monday, sunday = datetime.now(), datetime.now()
    monday = (monday - timedelta(days=(monday.weekday() - 7 * delta))).strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月',
                                                                                                          d='日')
    sunday = (sunday + timedelta(days=(6 - sunday.weekday() - 7 * -delta))).strftime('%Y{y}%m{m}%d{d}').format(y='年',
                                                                                                               m='月',
                                                                                                               d='日')
    return monday, sunday


def generate(name="巩光乾", last_week=False, ):
    monday, sunday = get_monday_sunday(-1 if last_week else 0)
    document = Document()

    document.add_heading('主要工作:', 0)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'{monday}--{sunday}')
    run.font.name = u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'{name}')
    run.font.name = u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    document.save(f'{monday}-{sunday} {name}.docx')


def main():
    Fire(generate)
